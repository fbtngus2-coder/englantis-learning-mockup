from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ENGLANTIS DO 활동 게임형 기획 및 이미지 에셋 요청서_중복정리본.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(8.8)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"

    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(width)
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, header, bold=True, color="1F4D78")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, (text, width) in enumerate(zip(row, widths)):
            cells[i].width = Inches(width)
            set_cell_text(cells[i], text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10)


def add_code_list(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)


def add_status_table(doc: Document, rows: list[list[str]]) -> None:
    add_table(
        doc,
        ["상태", "에셋/역할", "기존 참고 또는 요청 방식"],
        rows,
        [1.1, 1.8, 3.6],
    )


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "List Bullet 2"]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10)

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_intro(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("ENGLANTIS DO 활동 게임형 기획 및 이미지 에셋 요청서")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(21)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("중복 이미지 정리본 - 없는 이미지만 새로 제작하고, 기존 캐릭터는 참고 이미지 요청 방식으로 유지")
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")

    doc.add_heading("0. GPT에게 먼저 전달할 핵심 지시", level=1)
    add_bullet(doc, "새 이미지를 만들기 전에 아래의 기존 참고 이미지가 있는지 사용자에게 요청하세요.")
    add_bullet(doc, "Jiho, PIX, Dr. Lexi, Milo, Fogmon처럼 이미 존재하는 캐릭터는 새로 디자인하지 말고 기존 이미지를 기준으로 같은 캐릭터의 새 포즈만 제작하세요.")
    add_bullet(doc, "문서에서 '기존 있음 - 사용자에게 참고 이미지 요청'이라고 표시된 항목은 새로 만들지 말고, 사용자가 이미지를 제공하면 그 이미지와 같은 캐릭터성/색감/의상을 유지하세요.")
    add_bullet(doc, "문서에서 '신규 제작 필요'라고 표시된 항목만 새로 생성하세요.")
    add_bullet(doc, "이미지 안에는 한국어/영어 텍스트를 넣지 마세요. 텍스트는 나중에 UI에서 별도로 얹습니다.")

    doc.add_heading("1. 기존 이미지 확인 요약", level=1)
    add_status_table(
        doc,
        [
            ["기존 있음", "Jiho 기본/준비/생각/포인트/걷기/캐스팅/응원", "이미지/02_characters/jiho/ 안에 다수 있음. 새 포즈가 필요하면 사용자에게 jiho-idle.png, jiho-ready.png 등을 참고 이미지로 요청."],
            ["기존 있음", "PIX 기본/힌트/걱정/축하/비행/포인트", "이미지/02_characters/pix/ 안에 다수 있음. 새 포즈가 필요하면 사용자에게 pix-idle.png, pix-hint.png 등을 요청."],
            ["기존 있음", "Dr. Lexi 기본/칭찬/가르치는 포즈", "이미지/02_characters/dr-lexi/ 안에 있음. 새 코치 포즈가 필요하면 사용자에게 dr-lexi-neutral.png 또는 dr-lexi-teaching.png 요청."],
            ["기존 있음", "Milo 기본/생각/걷기/포인트/환영/축하", "이미지/02_characters/milo/ 안에 있음. 헷갈림/기쁨 같은 변형은 기존 Milo 참고 이미지 요청 후 제작."],
            ["기존 있음", "Fogmon 여러 종류와 상태", "이미지/02_characters/fogmon/variants/ 안에 vocab, sound, puzzle, lumio, shadow 계열 있음. 새 안개몬이 아니라면 기존 활용."],
            ["기존 있음", "주요 배경", "이미지/01_backgrounds/ 안에 battle street, rune chamber, rhythm stage, slingshot yard, magic classroom, library, plaza 등 있음."],
            ["기존 있음", "공통 효과", "이미지/03_game-elements/effects/, magic-runes/, projectiles/, ui/ 안에 마법진, 연기, 반짝임, 충격, 발사 궤적, 성공/실패 UI 있음."],
            ["기존 있음", "사이비 신도 활동 캐릭터", "이미지/eie_cultist_do_activity_assets_transparent/... 안에 로브 NPC, 신도, 무고한 NPC, 뿅망치 공격 포즈 있음."],
        ],
    )


ACTIVITIES = [
    {
        "title": "01. Sentence Purify - 안개 문장 정화",
        "story": "Grammia 광장의 문장석이 안개에 오염되어 핵심 단어가 사라졌다. Jiho는 PIX의 힌트를 받아 빈칸에 맞는 표현을 선택하고, 정화 주문으로 문장을 복구한다.",
        "flow": ["오염된 문장석 등장", "PIX가 뜻/상황 힌트 제공", "선택지 중 정답 선택", "정답이면 안개가 걷히고 문장이 빛남", "오답이면 안개가 짙어지고 추가 힌트 제공"],
        "existing": [
            ["기존 있음", "배경", "이미지/01_backgrounds/bg-fog-battle-street-night.png 또는 bg-fog-plaza-battle-night.png 사용 가능"],
            ["기존 있음", "Jiho / PIX", "기존 캐릭터 참고 이미지 요청 후 같은 캐릭터로 사용"],
            ["기존 있음", "Fogmon", "fogmon-grammia-puzzle, fogmon-mist-vocab 계열 사용 가능"],
            ["기존 있음", "공통 효과", "magic-circle-gold-ground, sparkle-cluster-gold, smoke-wisp-purple 활용 가능"],
        ],
        "new": [
            "sentence-purify/objects/obj_sentence_stone_blank.png - 텍스트 없는 오염된 문장석",
            "sentence-purify/objects/obj_sentence_stone_purified.png - 텍스트 없는 정화된 문장석",
            "sentence-purify/effects/fx_sentence_blank_glow.png - 빈칸 강조 빛",
            "sentence-purify/effects/fx_purify_beam.png - 선택지에서 문장석으로 날아가는 정화 빔",
        ],
        "ask": "Jiho와 PIX는 이미 있으니 새로 만들지 말고, 필요하면 사용자에게 기존 Jiho/PIX 참고 이미지를 요청한 뒤 같은 스타일의 정화 주문 포즈만 제작하세요.",
    },
    {
        "title": "02. Rune Puzzle - 문장 룬 조립",
        "story": "고대 룬 문이 문장 순서를 잃어 잠겨 있다. 단어 룬을 올바른 어순으로 꽂아야 문이 열리고 다음 구역으로 이동할 수 있다.",
        "flow": ["잠긴 룬 문 등장", "빈 슬롯과 단어 룬 카드 표시", "단어를 순서대로 선택", "맞는 순서면 룬이 결합", "문장이 완성되면 문이 열림"],
        "existing": [
            ["기존 있음", "배경", "이미지/01_backgrounds/bg-rune-puzzle-chamber.png 사용 가능"],
            ["기존 있음", "룬/마법진 효과", "이미지/03_game-elements/magic-runes/의 sigil, orb, magic-circle 활용 가능"],
            ["기존 있음", "Jiho / PIX", "기존 참고 이미지 요청 후 유지"],
        ],
        "new": [
            "rune-puzzle/objects/obj_rune_door_locked.png - 잠긴 고대 문",
            "rune-puzzle/objects/obj_rune_door_open.png - 열린 고대 문",
            "rune-puzzle/objects/obj_rune_slot_empty.png - 빈 룬 슬롯",
            "rune-puzzle/objects/obj_rune_word_tile.png - 텍스트 없는 단어 룬 타일",
            "rune-puzzle/effects/fx_rune_snap.png - 룬이 슬롯에 꽂히는 효과",
            "rune-puzzle/effects/fx_door_open_light.png - 문 열림 빛 효과",
        ],
        "ask": "단어는 이미지에 넣지 않습니다. 룬 타일은 비어 있는 판 형태로만 만들어주세요.",
    },
    {
        "title": "03. NPC Mission Talk - 주민 안내 대화",
        "story": "길을 잃은 주민 Milo가 중요한 장소를 찾지 못하고 있다. Jiho는 오늘 배운 표현을 사용해 주민에게 길을 알려주고, 도시의 길 표지판을 복원한다.",
        "flow": ["NPC가 문제 상황을 말함", "장소 단서 확인", "알맞은 응답 선택/입력", "정답이면 길이 빛남", "NPC가 목적지로 이동"],
        "existing": [
            ["기존 있음", "배경", "이미지/01_backgrounds/bg-npc-bakery-street-twilight.png 또는 bg-lingua-vita-plaza-day.png 사용 가능"],
            ["기존 있음", "Milo", "이미지/02_characters/milo/ 안에 기본/생각/걷기/포인트/축하 있음"],
            ["기존 있음", "PIX", "힌트/걱정/축하 포즈 있음"],
        ],
        "new": [
            "npc-mission-talk/objects/obj_route_signpost.png - 텍스트 없는 판타지 길 표지판",
            "npc-mission-talk/objects/obj_map_fragment.png - 작은 지도 조각",
            "npc-mission-talk/effects/fx_route_light_path.png - 바닥에 생기는 정답 길 빛",
            "npc-mission-talk/effects/fx_npc_understood_spark.png - NPC 이해 완료 효과",
            "npc-mission-talk/characters/npc_milo_confused.png - 기존 Milo 참고 후 헷갈림 표정이 꼭 필요할 때만 제작",
        ],
        "ask": "Milo는 이미 있으므로 새 캐릭터로 디자인하지 말고, 사용자에게 기존 Milo 이미지를 요청한 뒤 같은 인물의 헷갈림/기쁨 변형만 제작하세요.",
    },
    {
        "title": "04. Find the Cultist - 사이비 신도를 찾아라",
        "story": "수상한 로브를 입은 세 NPC가 같은 광장에 나타났다. 겉모습은 비슷하지만 한 명은 문법 신호가 깨진 문장을 말한다.",
        "flow": ["로브 NPC A/B/C 등장", "각 NPC의 문장 확인", "세 문장을 모두 들은 뒤 후보 지목", "뿅망치 판정", "정답이면 로브 해제/도망", "오답이면 무고한 NPC 반응"],
        "existing": [
            ["기존 있음", "로브 NPC/신도/무고한 NPC/뿅망치 공격 포즈", "이미지/eie_cultist_do_activity_assets_transparent/... 안에 01~10 PNG 있음"],
            ["기존 있음", "배경 대체", "bg-fog-battle-street-night.png 또는 bg-npc-bakery-street-twilight.png 사용 가능"],
            ["기존 있음", "연기/충격 효과", "impact-dust-gray, smoke-wisp-purple, sparkle-star-gold 활용 가능"],
        ],
        "new": [
            "find-the-cultist/backgrounds/bg_cultist_square_night.png - 필요 시 전용 수상한 광장 배경",
            "find-the-cultist/objects/obj_squeaky_hammer.png - 뿅망치 단독 오브젝트",
            "find-the-cultist/effects/fx_hammer_impact.png - 뿅망치 충격 별/먼지",
            "find-the-cultist/effects/fx_robe_reveal_smoke.png - 로브 해제 연기",
        ],
        "ask": "로브 NPC와 신도는 이미 있으므로 새로 만들지 마세요. 다른 포즈가 필요하면 사용자에게 기존 01~10 이미지를 참고 이미지로 요청하세요.",
    },
    {
        "title": "05. Fog Creature Battle Lab - 안개몬 배틀",
        "story": "안개몬은 악한 존재가 아니라 영어 신호가 뒤틀려 폭주한 존재다. 올바른 영어 신호로 반전 LV를 낮추고 정화한다.",
        "flow": ["도시/안개몬/배틀 유형 선택", "전투 화면 진입", "안개몬이 문제 제시", "선택/타이핑/말하기/듣기로 응답", "정답이면 공격/정화", "HP 또는 Fog Density 0이면 승리"],
        "existing": [
            ["기존 있음", "안개몬", "이미지/02_characters/fogmon/variants/에 mist-vocab, speakia-sound, grammia-puzzle, lumio-cloud, shadow-horned 있음"],
            ["기존 있음", "안개몬 상태", "idle, hit, purified, portrait 감정 상태 다수 있음"],
            ["기존 있음", "배경", "bg-fog-battle-street-night.png, bg-fog-plaza-battle-night.png 사용 가능"],
            ["기존 있음", "전투 효과", "slash-arc-gold, projectile-trail-gold-curved, magic-circle-gold-ground 사용 가능"],
        ],
        "new": [
            "fog-creature-battle-lab/ui/ui_battle_command_panel.png - 이미지 UI가 꼭 필요할 때만 제작. CSS로 대체 가능",
            "fog-creature-battle-lab/ui/ui_turn_counter.png - 턴 수 표시 프레임. CSS로 대체 가능",
            "fog-creature-battle-lab/ui/ui_enemy_hp_frame.png - 적 HP 프레임. CSS로 대체 가능",
            "fog-creature-battle-lab/effects/fx_shield_barrier.png - 방어막 효과",
            "fog-creature-battle-lab/fogmons/boss_fogmon_custom.png - 기존 shadow-horned로 부족할 경우에만 신규 보스 제작",
        ],
        "ask": "안개몬은 이미 종류가 많으므로 먼저 사용자에게 기존 fogmon 이미지를 요청하세요. 새 보스가 필요한 경우에만 기존 안개몬 톤을 참고해 제작하세요.",
    },
    {
        "title": "06. Exploration Expedition - 동굴/유적 탐험",
        "story": "동굴 깊은 곳에 잃어버린 영어 보물이 잠들어 있다. Jiho는 PIX의 힌트를 따라 알맞은 단어 유물을 골라가며 안쪽으로 들어간다.",
        "flow": ["동굴 입구에서 시작", "갈림길마다 힌트 제공", "알맞은 유물/단어/그림 선택", "정답이면 횃불 점등", "오답이면 안개 함정", "보물방에서 상자 획득"],
        "existing": [
            ["기존 있음", "대체 배경", "bg-explore-ruins-twilight.png 사용 가능"],
            ["기존 있음", "Jiho 걷기", "jiho-walk.png, jiho-walk-side.png 있음"],
            ["기존 있음", "PIX 안내", "pix-guide-point, pix-hint 사용 가능"],
        ],
        "new": [
            "exploration-expedition/backgrounds/bg_cave_entrance.png - 동굴 입구",
            "exploration-expedition/backgrounds/bg_cave_inner_path.png - 동굴 내부 길",
            "exploration-expedition/backgrounds/bg_cave_split_path.png - 갈림길",
            "exploration-expedition/backgrounds/bg_cave_treasure_room.png - 보물방",
            "exploration-expedition/objects/obj_word_artifact_correct.png - 정답 유물",
            "exploration-expedition/objects/obj_word_artifact_wrong.png - 함정 유물",
            "exploration-expedition/objects/obj_torch_off.png",
            "exploration-expedition/objects/obj_torch_on.png",
            "exploration-expedition/objects/obj_treasure_chest_closed.png",
            "exploration-expedition/objects/obj_treasure_chest_open.png",
            "exploration-expedition/objects/obj_trap_fog_tile.png",
            "exploration-expedition/effects/fx_treasure_glow.png",
        ],
        "ask": "Jiho/Pix는 기존 이미지를 사용합니다. 새 동굴 탐험 포즈가 필요하면 사용자에게 Jiho 기존 이미지를 요청한 뒤 같은 캐릭터로 제작하세요.",
    },
    {
        "title": "07. Magic Workshop - 마법 공방 제작",
        "story": "마법 공방의 장치가 고장났다. 제작 순서를 영어로 올바르게 배열해야 물약, 나침반, 열쇠 같은 마법 아이템이 완성된다.",
        "flow": ["제작 의뢰 등장", "재료와 순서 카드 제공", "순서 선택", "맞으면 제작 게이지 상승", "완성 아이템 획득", "틀리면 연기와 힌트"],
        "existing": [
            ["기존 있음", "대체 배경", "bg-magic-classroom-warm.png 사용 가능"],
            ["기존 있음", "Dr. Lexi", "dr-lexi-teaching.png, dr-lexi-neutral.png 사용 가능"],
            ["기존 있음", "공통 효과", "sparkle, smoke, magic circle 계열 활용 가능"],
        ],
        "new": [
            "magic-workshop/backgrounds/bg_magic_workshop_table.png - 공방 테이블 중심 배경",
            "magic-workshop/characters/npc_artisan_idle.png - 장인 NPC",
            "magic-workshop/characters/npc_artisan_success.png",
            "magic-workshop/characters/npc_artisan_worried.png",
            "magic-workshop/objects/obj_cauldron_idle.png",
            "magic-workshop/objects/obj_cauldron_success.png",
            "magic-workshop/objects/obj_cauldron_fail_smoke.png",
            "magic-workshop/objects/obj_recipe_scroll.png",
            "magic-workshop/objects/obj_ingredient_card.png - 텍스트 없는 재료 카드",
            "magic-workshop/objects/obj_magic_compass_complete.png",
            "magic-workshop/objects/obj_magic_key_complete.png",
            "magic-workshop/effects/fx_crafting_success_burst.png",
        ],
        "ask": "Dr. Lexi를 공방 코치로 쓸 경우 새 캐릭터를 만들지 말고 기존 Dr. Lexi 이미지를 사용자에게 요청하세요. 장인 NPC는 신규 제작 가능합니다.",
    },
    {
        "title": "08. Detective Case - 탐정 사건 해결",
        "story": "Readia 기록 보관소에 이상한 사건이 발생했다. 학습자는 증언과 기록을 읽고 사건 보드에 결정적인 증거를 연결해야 한다.",
        "flow": ["사건 의뢰 등장", "짧은 기록/증언 제시", "증거 카드 선택", "정답이면 사건 보드에 선 연결", "모든 단서를 연결하면 사건 해결"],
        "existing": [
            ["기존 있음", "대체 배경", "bg-library-teach-room.png 또는 bg-master-study-room.png 사용 가능"],
            ["기존 있음", "공통 종이/마법 효과 일부", "story_intro/object, sparkle 계열 활용 가능"],
        ],
        "new": [
            "detective-case/backgrounds/bg_archive_detective_room.png - 사건 보드가 어울리는 기록 보관소",
            "detective-case/characters/npc_archivist_miro_idle.png",
            "detective-case/characters/npc_archivist_miro_point.png",
            "detective-case/characters/npc_archivist_miro_surprised.png",
            "detective-case/objects/obj_case_board_empty.png",
            "detective-case/objects/obj_case_board_connected.png",
            "detective-case/objects/obj_evidence_card.png - 텍스트 없는 증거 카드",
            "detective-case/objects/obj_evidence_card_selected.png",
            "detective-case/objects/obj_evidence_card_wrong.png",
            "detective-case/objects/obj_detective_stamp_solved.png",
            "detective-case/effects/fx_clue_line_glow.png",
        ],
        "ask": "Archivist Miro가 기존 Milo와 다른 캐릭터라면 신규 제작. 기존 Milo를 변형할 의도라면 사용자에게 Milo 참고 이미지를 요청하세요.",
    },
    {
        "title": "09. Rune Runner - WASD 유적 달리기",
        "story": "무너지는 유적에서 영어 룬이 흩어졌다. Jiho는 제한 시간 안에 정답 룬을 모으고 오답 안개와 함정을 피해 탈출해야 한다.",
        "flow": ["던전 맵 시작", "유물과 함정 배치", "WASD/방향키 이동", "정답 유물 수집", "함정 접촉 시 안개 증가", "목표 수집 후 탈출문 개방"],
        "existing": [
            ["기존 있음", "대체 배경", "bg-explore-ruins-twilight.png, bg-rune-puzzle-chamber.png 사용 가능"],
            ["기존 있음", "Jiho 걷기", "jiho-walk.png, jiho-walk-side.png 있음"],
            ["기존 있음", "룬/마법 효과", "magic-runes, sparkle, speed-ring-gold 활용 가능"],
        ],
        "new": [
            "rune-runner/backgrounds/bg_runner_dungeon_map.png - 탑다운/사이드뷰용 던전 맵",
            "rune-runner/objects/obj_collectible_rune_correct.png",
            "rune-runner/objects/obj_collectible_rune_wrong.png",
            "rune-runner/objects/obj_spike_trap.png",
            "rune-runner/objects/obj_fog_hazard.png",
            "rune-runner/objects/obj_exit_portal_closed.png",
            "rune-runner/objects/obj_exit_portal_open.png",
            "rune-runner/effects/fx_collect_spark.png",
            "rune-runner/effects/fx_trap_hit_smoke.png",
        ],
        "ask": "Jiho runner 포즈는 기존 jiho-walk 계열이 있으므로 먼저 사용자에게 해당 이미지를 요청하세요. 부족할 경우에만 같은 캐릭터의 달리기 포즈를 제작하세요.",
    },
    {
        "title": "10. Rhythm Spell - 리듬 주문",
        "story": "Speakia 무대의 주문진은 영어 리듬에 반응한다. 학습자는 박자에 맞춰 룬 키를 눌러 봉인된 목소리를 되살린다.",
        "flow": ["리듬 주문진 등장", "노트가 판정선으로 이동", "타이밍에 맞춰 입력", "Perfect/Good/Miss 판정", "곡 종료 후 정확도 등급 표시"],
        "existing": [
            ["기존 있음", "배경", "bg-rhythm-spell-stage.png 사용 가능"],
            ["기존 있음", "Jiho 캐스팅", "jiho-casting.png 사용 가능"],
            ["기존 있음", "음파/마법 효과", "wind-trail-blue, sparkle, magic-circle 계열 활용 가능"],
        ],
        "new": [
            "rhythm-spell/objects/obj_beat_note_a.png - 텍스트 없는 리듬 노트",
            "rhythm-spell/objects/obj_beat_note_s.png",
            "rhythm-spell/objects/obj_beat_note_d.png",
            "rhythm-spell/objects/obj_beat_note_f.png",
            "rhythm-spell/objects/obj_beat_note_miss.png",
            "rhythm-spell/ui/ui_rhythm_judgement_line.png",
            "rhythm-spell/ui/ui_combo_badge.png - 텍스트 없는 콤보 배지",
            "rhythm-spell/effects/fx_perfect_hit.png",
            "rhythm-spell/effects/fx_good_hit.png",
            "rhythm-spell/effects/fx_miss_shatter.png",
            "rhythm-spell/effects/fx_spell_circle_complete.png",
        ],
        "ask": "Jiho는 기존 캐스팅 포즈가 있으므로 새 캐릭터를 만들지 마세요. 리듬 전용 포즈가 필요하면 기존 Jiho 이미지를 사용자에게 요청하세요.",
    },
    {
        "title": "11. Slingshot Academy - 새총 사격",
        "story": "Vovania 사격장에서 떠다니는 단어 표적들이 안개에 섞여 있다. Jiho는 마법 새총으로 정답 표적만 맞혀야 한다.",
        "flow": ["표적 여러 개 등장", "문제 힌트 제시", "새총을 당겨 조준", "놓으면 마법 구슬 발사", "정답 표적 폭발", "오답 표적은 튕김"],
        "existing": [
            ["기존 있음", "배경", "bg-slingshot-training-yard.png 사용 가능"],
            ["기존 있음", "표적/발사 궤적 일부", "target-ring-cyan-flat, projectile-trail-gold-curved 사용 가능"],
            ["기존 있음", "Jiho 기본", "jiho-ready, jiho-point 등 참고 가능"],
        ],
        "new": [
            "slingshot-academy/characters/jiho_slingshot_aim.png - 기존 Jiho 참고 후 제작",
            "slingshot-academy/characters/jiho_slingshot_release.png - 기존 Jiho 참고 후 제작",
            "slingshot-academy/characters/jiho_slingshot_success.png - 기존 Jiho 참고 후 제작",
            "slingshot-academy/objects/obj_slingshot.png - 새총 단독",
            "slingshot-academy/objects/obj_magic_stone_projectile.png",
            "slingshot-academy/objects/obj_target_correct.png - 텍스트 없는 정답 표적",
            "slingshot-academy/objects/obj_target_wrong.png - 텍스트 없는 오답 표적",
            "slingshot-academy/objects/obj_target_hit.png",
            "slingshot-academy/objects/obj_target_shield_block.png",
            "slingshot-academy/effects/fx_target_burst.png",
            "slingshot-academy/effects/fx_wrong_target_bounce.png",
            "slingshot-academy/effects/fx_combo_spark.png",
        ],
        "ask": "새총 포즈는 Jiho 캐릭터 유지가 중요합니다. GPT는 사용자에게 기존 Jiho 참고 이미지를 요청한 뒤 같은 얼굴/의상/색감으로 조준/발사/성공 포즈를 만들어야 합니다.",
    },
]


def add_activity(doc: Document, item: dict) -> None:
    doc.add_heading(item["title"], level=1)
    p = doc.add_paragraph()
    p.add_run("스토리: ").bold = True
    p.add_run(item["story"])

    doc.add_heading("플레이 흐름", level=2)
    for step in item["flow"]:
        add_bullet(doc, step)

    doc.add_heading("기존 이미지 / 참고 요청", level=2)
    add_status_table(doc, item["existing"])

    doc.add_heading("새로 제작할 이미지", level=2)
    add_code_list(doc, item["new"])

    doc.add_heading("GPT에게 요청할 말", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(item["ask"])
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10)


def add_appendix(doc: Document) -> None:
    doc.add_heading("부록 A. 이미지 제작 공통 규칙", level=1)
    add_bullet(doc, "배경은 16:9 비율. 전경 캐릭터 없이 UI가 올라갈 빈 공간을 남긴다.")
    add_bullet(doc, "캐릭터/오브젝트/효과는 투명 배경 PNG를 우선한다.")
    add_bullet(doc, "이미지 안에 텍스트를 넣지 않는다.")
    add_bullet(doc, "기존 캐릭터가 있는 경우, 반드시 사용자에게 기존 이미지를 요청하고 같은 캐릭터로 변형한다.")
    add_bullet(doc, "파일명은 문서에 적힌 경로와 이름을 우선 사용한다.")

    doc.add_heading("부록 B. GPT가 사용자에게 물어봐야 하는 예시 문장", level=1)
    examples = [
        "Jiho는 기존 이미지가 있는 캐릭터입니다. 같은 캐릭터의 새 포즈를 만들기 위해 jiho-idle.png와 jiho-ready.png를 업로드해 주세요.",
        "PIX는 기존 이미지가 있으므로 새로 디자인하지 않겠습니다. pix-idle.png와 pix-hint.png를 참고 이미지로 주시면 같은 톤으로 새 표정을 만들겠습니다.",
        "Fogmon은 기존 variants 폴더에 여러 종류가 있습니다. 새 보스를 만들지, 기존 shadow-horned 계열을 보스로 쓸지 결정해 주세요.",
        "이 배경은 기존 bg-rune-puzzle-chamber.png로 대체 가능해 보입니다. 새 배경이 꼭 필요한지 확인해 주세요.",
    ]
    for example in examples:
        add_bullet(doc, example)

    doc.add_heading("부록 C. 신규 제작 우선순위", level=1)
    add_status_table(
        doc,
        [
            ["1순위", "오늘 과정 DO 4종", "Sentence Purify, Rune Puzzle, NPC Mission Talk, Find the Cultist의 부족 오브젝트/효과부터 제작"],
            ["2순위", "조작감 강한 활동", "Rune Runner, Rhythm Spell, Slingshot Academy의 조작용 오브젝트와 효과 제작"],
            ["3순위", "확장형 활동", "Exploration, Workshop, Detective, Battle Lab의 전용 배경과 NPC 제작"],
        ],
    )


def main() -> None:
    doc = Document()
    configure_doc(doc)
    add_intro(doc)

    for idx, activity in enumerate(ACTIVITIES):
        if idx in {4, 8}:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_activity(doc, activity)

    add_appendix(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
